import os
import io
import sys
import hashlib
from docx import Document
import PyPDF2
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

_form_endpoint = os.getenv("AZURE_FORM_RECOGNIZER_ENDPOINT")
_form_key = os.getenv("AZURE_FORM_RECOGNIZER_KEY")
_openai_key = os.getenv("AZURE_OPENAI_API_KEY")
_openai_base_url = os.getenv("AZURE_OPENAI_BASE_URL")
_openai_deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")

# A PDF with fewer than this many characters per page (locally) is treated as
# scanned/image-only and routed to Azure OCR.
_MIN_CHARS_PER_PAGE = 100

_OCR_CACHE_DIR = os.path.join(os.path.dirname(__file__), ".ocr_cache")


# -- Local extractors (free, no API) -------------------------------------------

def extract_text_from_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", "replace")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract paragraphs AND table cells (UoA templates put clauses in tables)."""
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    return "\n".join(p.extract_text() for p in reader.pages if p.extract_text())


def _pdf_page_count(file_bytes: bytes) -> int:
    try:
        return max(1, len(PyPDF2.PdfReader(io.BytesIO(file_bytes)).pages))
    except Exception:
        return 1


# -- OCR cache -----------------------------------------------------------------

def _ocr_cache_path(file_bytes: bytes) -> str:
    digest = hashlib.sha256(file_bytes).hexdigest()[:32]
    return os.path.join(_OCR_CACHE_DIR, digest + ".txt")


def _read_ocr_cache(file_bytes: bytes):
    path = _ocr_cache_path(file_bytes)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as fh:
            return fh.read()
    return None


def _write_ocr_cache(file_bytes: bytes, text: str) -> None:
    os.makedirs(_OCR_CACHE_DIR, exist_ok=True)
    with open(_ocr_cache_path(file_bytes), "w", encoding="utf-8") as fh:
        fh.write(text)


# -- Azure OCR (paid; called only as a fallback) -------------------------------

def extract_text_with_azure(file_bytes: bytes, source_name: str = "uploaded document") -> str:
    """OCR a document with Azure Document Intelligence (prebuilt-read).

    Results are cached by content hash so the same file is never OCR'd twice.
    A guard line is printed before any real Azure call so charges are visible.
    """
    cached = _read_ocr_cache(file_bytes)
    if cached is not None:
        print(f"[Azure OCR] cache hit for '{source_name}' - no API call.", file=sys.stderr)
        return cached

    print(
        f"[Azure OCR] >>> Calling Azure Document Intelligence (prebuilt-read) for "
        f"'{source_name}' ({len(file_bytes):,} bytes). THIS INCURS COST. <<<",
        file=sys.stderr,
    )
    client = DocumentAnalysisClient(_form_endpoint, AzureKeyCredential(_form_key))
    poller = client.begin_analyze_document("prebuilt-read", file_bytes)
    result = poller.result()
    text = result.content or ""
    n_pages = len(result.pages) if getattr(result, "pages", None) else "?"
    print(
        f"[Azure OCR] done: '{source_name}' -> {len(text):,} chars from {n_pages} page(s).",
        file=sys.stderr,
    )
    _write_ocr_cache(file_bytes, text)
    return text


# -- Dispatcher: local-first, Azure OCR only for scanned PDFs -------------------

def extract_text(file_bytes: bytes, filename: str = "") -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return extract_text_from_txt(file_bytes)
    if ext == ".docx":
        return extract_text_from_docx(file_bytes)

    # PDF (or unknown): try local first, fall back to OCR if text is sparse.
    try:
        local = extract_text_from_pdf(file_bytes)
    except Exception:
        local = ""

    pages = _pdf_page_count(file_bytes)
    if local.strip() and len(local.strip()) >= _MIN_CHARS_PER_PAGE * pages:
        return local

    return extract_text_with_azure(file_bytes, source_name=filename or "PDF")


def analyze_with_foundry_model(text: str) -> str:
    client = OpenAI(api_key=_openai_key, base_url=_openai_base_url)
    response = client.responses.create(
        model=_openai_deployment,
        input=f"Analyze this contract text: {text[:16000]}"
    )
    return response.output_text
